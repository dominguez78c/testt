import os
import json
import zipfile
import io
import shutil 
import base64
from flask import Flask, request, jsonify, send_file, render_template 
from flask_cors import CORS 
from werkzeug.utils import secure_filename

# LIBRERÍAS DE DIBUJO Y PLOTEO
import drawsvg as dw                    # Para generar el archivo SVG final (Vectorial)
import matplotlib.pyplot as plt         # Para generar el PNG de previsualización (Matriz de puntos)
from matplotlib.patches import PathPatch
from matplotlib.path import Path as MplPath 

# LIBRERÍAS PARA WORD (DOCX)
from docx import Document 
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Las importaciones OxmlElement y qn fueron removidas ya que no se usan en la versión actual.

# LIBRERÍA DE GOOGLE AI
from google import genai
from google.genai import types

# ==========================================================
# 🔑 CONFIGURACIÓN DE LA CLAVE API 🔑
# ==========================================================
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY") 

# ==========================================================
# CONFIGURACIÓN Y RUTAS 
# ==========================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'outputs')
TEMP_FOLDER = os.path.join(OUTPUT_FOLDER, 'temp_images') 

app = Flask(__name__)
CORS(app) 
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['TEMP_FOLDER'] = TEMP_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(TEMP_FOLDER, exist_ok=True)

# ==========================================================
# RUTA PRINCIPAL
# ==========================================================
@app.route('/')
def index():
    # Asume que 'index.html' existe en la carpeta 'templates'
    return render_template('index.html')



## Funciones Auxiliares de Archivo

def guardar_path_como_png(path_data: str, w_draw: float, h_draw: float, output_path: str):
    """
    Genera una previsualización PNG del dibujo de la pieza usando Matplotlib.
    Si es una forma compleja, se usa un recuadro de referencia.
    """
    # Definimos el tamaño de la figura (constante para el DOCX)
    fig_size = 4
    fig, ax = plt.subplots(figsize=(fig_size, fig_size))
    
    # Configuramos el límite del plot (con margen)
    margin = 0.1 * max(w_draw, h_draw)
    ax.set_xlim(-margin, w_draw + margin)
    ax.set_ylim(-margin, h_draw + margin)
    ax.set_aspect('equal', adjustable='box')
    ax.axis('off') # Ocultar ejes

    # Se usa una línea discontinua como caja delimitadora para formas irregulares
    if path_data and len(str(path_data).strip()) > 10:
        # Forma irregular: Dibujamos la caja delimitadora y una etiqueta de advertencia
        ax.add_patch(plt.Rectangle((0, 0), w_draw, h_draw, 
                                   fill=False, edgecolor='red', linewidth=2, linestyle='--'))
        ax.text(w_draw / 2, h_draw / 2, 'FORMA IRREGULAR (Ver SVG)', 
                ha='center', va='center', fontsize=8, color='red')
    else:
        # Rectángulo estándar
        ax.add_patch(plt.Rectangle((0, 0), w_draw, h_draw, 
                                   fill=False, edgecolor='black', linewidth=2))
        ax.text(w_draw / 2, h_draw / 2, 'RECTANGULAR', 
                ha='center', va='center', fontsize=10, color='black')

    # Guardar y cerrar la figura
    plt.savefig(output_path, dpi=150, bbox_inches='tight', pad_inches=0.1)
    plt.close(fig) 
    return output_path



## 1. Análisis del Render (Gemini)

def analizar_y_generar_desglose(data: dict, path_render: str):
    """
    Función que usa Gemini para obtener el desglose y el código SVG Path.
    """
    print(f"\n[IA] Conectando a Gemini para analizar formas y desglose...")
    
    if not GEMINI_API_KEY:
        print("[ERROR CRÍTICO] Falta GEMINI_API_KEY. Usando fallback.")
        return json.loads("""[{"nombre": "Error Config", "material": "N/A", "dimensiones_cm": [10,10,10], "cantidad_unidades": 0, "notas_fabricacion": "Configure la API KEY", "svg_path_d": null}]""")

    uploaded_file = None
    try:
        client = genai.Client(api_key=GEMINI_API_KEY)
        uploaded_file = client.files.upload(file=path_render)

        system_prompt = (
            "Eres un ingeniero de producción experto en CAD/CNC. Analiza el render y genera un BOM (JSON). "
            "Incluye 'nombre', 'material', 'dimensiones_cm' [largo, alto, espesor], 'cantidad_unidades', 'notas_fabricacion'. "
            "CRÍTICO - GEOMETRÍA: Se debe incluir el campo 'svg_path_d'. Si la pieza es RECTANGULAR/CUADRADA, 'svg_path_d' debe ser null. "
            "Si la pieza es IRREGULAR (nube, arco, letra, silueta), DEBES generar un string con el atributo 'd' de un SVG path "
            "que represente esa forma exacta. Asume que el path se dibuja en un lienzo de 100x100 unidades (será escalado)."
        )

        user_prompt = (
            f"Analiza este render. Dimensiones totales: {data['frente']}x{data['fondo']}x{data['altura']} cm. "
            f"Contexto: {data['prompt']}. Genera el desglose y el código de trazado 'svg_path_d' si es necesario."
        )

        response = client.models.generate_content(
            model='gemini-2.0-flash', 
            contents=[uploaded_file, user_prompt],
            config=types.GenerateContentConfig(
                system_instruction=system_prompt,
                response_mime_type="application/json"
            )
        )
        
        return json.loads(response.text)

    except Exception as e:
        print(f"[ERROR GEMINI] Fallo al generar el contenido: {e}")
        return json.loads("""[{"nombre": "Muro Base (Fallback)", "material": "MDF", "dimensiones_cm": [100, 200, 5], "cantidad_unidades": 1, "notas_fabricacion": "Generado por error de conexión IA", "svg_path_d": null}]""")
        
    finally:
        if uploaded_file:
            try:
                client = genai.Client(api_key=GEMINI_API_KEY)
                client.files.delete(name=uploaded_file.name)
            except Exception as cleanup_e:
                pass



## 2. Generación de Planos SVG y PNG

def generar_planos_vectoriales_svg(componentes: list):
    """
    Genera archivos SVG (corte CNC) y las previsualizaciones PNG (DOCX).
    """
    svg_folder = os.path.join(app.config['OUTPUT_FOLDER'], "planos_svg")
    if os.path.exists(svg_folder): shutil.rmtree(svg_folder)
    os.makedirs(svg_folder, exist_ok=True)
    
    png_folder = app.config['TEMP_FOLDER'] 
    
    svg_files = []
    temp_png_files = []
    
    for comp in componentes:
        l, a, _ = comp['dimensiones_cm']
        
        # 1cm = 10 unidades de dibujo
        w_draw = l * 10
        h_draw = a * 10
        
        path_data = comp.get('svg_path_d')

        # --- 2A: Generación del archivo SVG para corte (Con drawsvg) ---
        canvas_width = w_draw + 100
        canvas_height = h_draw + 100
        dwg = dw.Drawing(canvas_width, canvas_height)
        
        if path_data and len(str(path_data).strip()) > 10:
            # DIBUJAR PATH COMPLEJO
            scale_x = w_draw / 100
            scale_y = h_draw / 100
            shape = dw.Path(d=str(path_data), fill='none', stroke='#E50000', stroke_width=4) 
            shape.args['transform'] = f"translate(50,50) scale({scale_x}, {scale_y})"
            dwg.append(shape)
            dwg.append(dw.Text(f"{comp['nombre']} (IRREGULAR)", 16, 50, 30))
        else:
            # DIBUJAR RECTÁNGULO
            dwg.append(dw.Rectangle(50, 50, w_draw, h_draw, fill='none', stroke='black', stroke_width=2))
            dwg.append(dw.Text(f"{comp['nombre']} (RECTANGULAR)", 16, 50, 30))
        
        dwg.append(dw.Text(f"Mat: {comp['material']} | {l}x{a}x{comp['dimensiones_cm'][2]}cm", 12, 50, canvas_height - 20))
        
        svg_fname = secure_filename(comp['nombre']) + '.svg'
        svg_path = os.path.join(svg_folder, svg_fname)
        dwg.save_svg(svg_path)
        svg_files.append(svg_path)
        
        # --- 2B: Generación del archivo PNG para el DOCX (Con Matplotlib) ---
        png_fname = svg_fname.replace('.svg', '.png')
        png_path = os.path.join(png_folder, png_fname)
        
        guardar_path_como_png(path_data, w_draw, h_draw, png_path)
        temp_png_files.append(png_path)
        comp['png_path'] = png_path 

    print(f"[SVG] Generados {len(svg_files)} archivos vectoriales.")
    print(f"[PNG] Generados {len(temp_png_files)} previsualizaciones para DOCX.")
    return svg_folder, svg_files, temp_png_files



## 3. Generación de Manual Word

def generar_manual_word(componentes: list, path_render: str, data: dict):
    doc_path = os.path.join(app.config['OUTPUT_FOLDER'], "Manual_Produccion.docx")
    document = Document()
    temp_png_files = [] 
    
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # TÍTULO Y RESUMEN
    document.add_heading("MANUAL DE PRODUCCIÓN TÉCNICA", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Proyecto: {data.get('prompt', 'Stand Custom')[:50]}...").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("\n" + "=" * 60 + "\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # SECCIÓN 1: VISTA GENERAL
    document.add_heading("1. Especificaciones Generales y Render", level=1)
    
    table_general = document.add_table(rows=1, cols=2)
    table_general.autofit = False 
    
    cell_img = table_general.cell(0, 0)
    cell_img.width = Inches(3.5)
    if os.path.exists(path_render):
        cell_img.paragraphs[0].add_run().add_picture(path_render, width=Inches(3.2))
    
    cell_data = table_general.cell(0, 1)
    cell_data.width = Inches(2.5)
    p_data = cell_data.paragraphs[0]
    p_data.add_run("DIMENSIONES TOTALES\n").bold = True
    p_data.add_run(f"Frente: {data['frente']} cm\n")
    p_data.add_run(f"Fondo: {data['fondo']} cm\n")
    p_data.add_run(f"Altura: {data['altura']} cm\n")

    document.add_paragraph("\n")

    # SECCIÓN 2: DESGLOSE DE PIEZAS (FICHA TÉCNICA)
    document.add_heading("2. Desglose Detallado (BOM)", level=1)
    
    for i, comp in enumerate(componentes):
        document.add_heading(f"Pieza {i+1}: {comp['nombre']}", level=2)
        
        t_pieza = document.add_table(rows=1, cols=2)
        t_pieza.style = 'Table Grid'
        
        # C1: DIBUJO 2D DEL CORTE (PNG generado)
        c_img = t_pieza.cell(0, 0)
        c_img.width = Inches(2.5)
        p_c_img = c_img.paragraphs[0]
        p_c_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        png_path = comp.get('png_path')
        
        if png_path and os.path.exists(png_path):
            p_c_img.add_run().add_picture(png_path, width=Inches(2.3)) 
            temp_png_files.append(png_path)
        else:
            p_c_img.add_run("[❌ Dibujo no disponible]")

        # C2: DATOS TÉCNICOS
        c_info = t_pieza.cell(0, 1)
        
        def add_field(paragraph, label, value):
            run = paragraph.add_run(f"{label}: ")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102) 
            paragraph.add_run(f"{value}\n")

        p_info = c_info.paragraphs[0]
        l, a, e = comp['dimensiones_cm']
        
        add_field(p_info, "MATERIAL", comp['material'])
        add_field(p_info, "MEDIDAS (cm)", f"{l} x {a} x {e}")
        add_field(p_info, "CANTIDAD", str(comp['cantidad_unidades']))
        
        if comp.get('notas_fabricacion'):
            p_note = c_info.add_paragraph()
            run_note = p_note.add_run(f"Nota: {comp['notas_fabricacion']}")
            run_note.italic = True
            run_note.font.size = Pt(9)
            
        svg_fname = secure_filename(comp['nombre']) + '.svg'
        c_info.add_paragraph(f"Archivo SVG: {svg_fname}")

        document.add_paragraph("\n") 

    document.save(doc_path)
    print(f"[DOCX] Manual profesional generado en: {doc_path}")
    return doc_path, temp_png_files


# ==========================================================
# 4. ENDPOINT DE FLASK PARA PROCESAMIENTO
# ==========================================================

@app.route('/generar', methods=['POST'])
def generar_manual():
    temp_files = []
    path_render = None
    doc_path = None
    svg_files = []
    
    try:
        # 1. Cargar datos
        file = request.files['render']
        filename = secure_filename(file.filename)
        path_render = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(path_render)
        data = json.loads(request.form['data'])
        
        # 2. Generación de Archivos
        componentes = analizar_y_generar_desglose(data, path_render)
        svg_folder, svg_files, temp_png_files = generar_planos_vectoriales_svg(componentes) 
        doc_path, docx_png_files = generar_manual_word(componentes, path_render, data)
        
        # 3. Compresión en ZIP
        memory_file = io.BytesIO()
        with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.write(doc_path, "Manual_Produccion.docx")
            zf.write(path_render, "Render_Original.jpg")
            
            for svg in svg_files:
                zf.write(svg, f"Planos_Corte/{os.path.basename(svg)}")
        
        memory_file.seek(0)
        
        # Archivos para limpieza
        temp_files = [path_render, doc_path] + svg_files + temp_png_files + docx_png_files
        
        return send_file(
            memory_file, 
            mimetype='application/zip', 
            as_attachment=True, 
            download_name='Proyecto_Produccion.zip'
        )

    except Exception as e:
        print(f"[ERROR FATAL] Error en la función generar_manual: {e}")
        return jsonify({"message": f"Error interno: {str(e)}"}), 500
        
    finally:
        # 4. Limpieza FINAL
        print("[CLEANUP] Iniciando limpieza de archivos temporales.")
        for f in temp_files:
            if os.path.exists(f): 
                try: os.remove(f)
                except: pass
        
        if os.path.exists(os.path.join(app.config['OUTPUT_FOLDER'], "planos_svg")):
            shutil.rmtree(os.path.join(app.config['OUTPUT_FOLDER'], "planos_svg"))
        if os.path.exists(app.config['TEMP_FOLDER']):
            shutil.rmtree(app.config['TEMP_FOLDER'])
        
        print("[CLEANUP] Limpieza finalizada.")


if __name__ == '__main__':
    if not GEMINI_API_KEY:
        print("\n=========================================================================")
        print("⚠️  ADVERTENCIA: GEMINI_API_KEY NO DETECTADA. La IA usará datos simulados.")
        print("=========================================================================\n")
    
    print("\n🚀 Servidor Flask iniciado. Visita: http://127.0.0.1:5000")
    print("---------------------------------------------------------")
    app.run(host='127.0.0.1', port=5000, debug=True)